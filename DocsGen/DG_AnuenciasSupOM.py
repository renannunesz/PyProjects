import os
import tkinter as tk
import comtypes.client
import re
from plyer import notification
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from ttkbootstrap import Style
from PIL import Image, ImageTk
from datetime import datetime

# Obter a data de hoje
hoje = datetime.today()
# Formatar a data como DDMMAAAA
data_hoje = hoje.strftime("%d%m%Y")

# Obtém o diretório onde o script está localizado
caminho_base = os.path.dirname(os.path.abspath(__file__))

class genAnuencias:
    def __init__(self, root):
        # Crie a interface gráfica
        style = Style(theme='darkly')
        self.root = root
        root = style.master
        root.title("DocsGen - Anuências - Supervisor de O&M")
        frame = ttk.Frame(root)
        frame['padding'] = (10, 10, 10, 10)
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        self.dir_pasta_selecionada = ""

        # Carregar a imagem
        self.bg_image = Image.open(r"C:\PyProjects\DocsGen\bkg_anuencias.png")
        self.bg_photo = ImageTk.PhotoImage(self.bg_image)
        
        # Inserir a imagem no topo do formulário
        ttk.Label(frame, image=self.bg_photo).grid(row=0, column=0, columnspan=6, pady=(0, 20))        

        #LabelFrame dentro do Frame - Dados Documento
        lblframe_dados_documentos = ttk.LabelFrame(frame, text="Dados Funcionário:", padding=10)
        lblframe_dados_documentos.grid(row=1, column=0, sticky="ew", padx=10, pady=5)   

        # Labels
        ttk.Label(lblframe_dados_documentos, text=" Funcionário: ").grid(row=0, column=0, sticky=tk.W)
        self.entr_funcionario = ttk.Entry(lblframe_dados_documentos, width=60)
        self.entr_funcionario.grid(row=0, column=1, sticky=tk.W, pady=10)

        ttk.Label(lblframe_dados_documentos, text=" CPF: ").grid(row=0, column=2, sticky=tk.W)
        self.entr_CPF = ttk.Entry(lblframe_dados_documentos, width=30)
        self.entr_CPF.grid(row=0, column=3, sticky=tk.W, pady=10)

        ttk.Label(lblframe_dados_documentos, text=" Iniciais: ").grid(row=0, column=4, sticky=tk.W)
        self.entr_apelido = ttk.Entry(lblframe_dados_documentos)
        self.entr_apelido.grid(row=0, column=5, sticky=tk.W, pady=10)

        #LabelFrame dentro do Frame - Anuencias
        lblframe_anuencias = ttk.LabelFrame(frame, text="Anuências:", padding=10)
        lblframe_anuencias.grid(row=2, column=0, sticky="ew", padx=10, pady=5)   

        ttk.Label(lblframe_anuencias, text="Supervisor de O&M: ").grid(row=0, column=0, padx=(0, 50), sticky=tk.W)

        # Variáveis para os Checkbuttons
        self.var_nr10_sup_oem = tk.BooleanVar(value=False)
        self.var_nr10_sep_sup_oem = tk.BooleanVar(value=False)
        self.var_nr12_sup_oem = tk.BooleanVar(value=False)
        self.var_nr33_sup_oem = tk.BooleanVar(value=False)
        self.var_nr35_sup_oem = tk.BooleanVar(value=False)    

        # Checkbox
        self.ckbx_nr10_sup_oem = ttk.Checkbutton(lblframe_anuencias, text="NR10", variable=self.var_nr10_sup_oem)
        self.ckbx_nr10_sup_oem.grid(row=0, column=1, sticky=tk.W, padx=(0, 10))
        self.ckbx_nr10_sep_sup_oem = ttk.Checkbutton(lblframe_anuencias, text="NR10 SEP", variable=self.var_nr10_sep_sup_oem).grid(row=0, column=2, sticky=tk.W, padx=(0, 10))
        self.ckbx_nr12_sup_oem = ttk.Checkbutton(lblframe_anuencias, text="NR12", variable=self.var_nr12_sup_oem).grid(row=0, column=3, sticky=tk.W, padx=(0, 10))
        self.ckbx_nr33_sup_oem = ttk.Checkbutton(lblframe_anuencias, text="NR33", variable=self.var_nr33_sup_oem).grid(row=0, column=4, sticky=tk.W, padx=(0, 10))
        self.ckbx_nr35_sup_oem = ttk.Checkbutton(lblframe_anuencias, text="NR35", variable=self.var_nr35_sup_oem).grid(row=0, column=5, sticky=tk.W, padx=(0, 10))

        #LabelFrame dentro do Frame - Opções
        lblframe_opcoes = ttk.LabelFrame(frame, text="Opções: ", padding=10)
        lblframe_opcoes.grid(row=3, column=0, sticky="ew", padx=10, pady=5)                

        # Btn  
        ttk.Label(lblframe_opcoes, text="Salvar Arquivos em: ").grid(row=0, column=0, sticky=tk.W)      
        self.btn_selecionar_pasta = ttk.Button(lblframe_opcoes, text="...", command=self.selecionar_pasta).grid(row=0, column=1, sticky=tk.W, padx=(0, 10))   
        self.lbl_pastaselecionada = tk.Label(lblframe_opcoes, text="Selecione a pasta", wraplength=350)
        self.lbl_pastaselecionada.grid(row=0, column=2, columnspan=4, sticky=tk.W, padx=(0,10))

        # Botão
        self.btn_gerar = ttk.Button(frame, text="Gerar Anuências", command=self.verificar_checkbuttons).grid(row=4, column=0, sticky=tk.W, padx=(0, 10), pady=10)

        # Defina os dados
        self.dados = {}       

    def selecionar_pasta(self):
        pasta_selecionada = filedialog.askdirectory()
        if pasta_selecionada:
            self.lbl_pastaselecionada.config(text=pasta_selecionada)   
            self.dir_pasta_selecionada = pasta_selecionada  

    def substituir_texto(self, paragraph, substituicoes):
        # Substituir o texto nos runs do parágrafo
        for palavra_antiga, palavra_nova in substituicoes.items():
            if palavra_antiga in paragraph.text:
                if palavra_antiga == 'NOMEFUNCIONARIO':
                    for run in paragraph.runs:
                        if palavra_antiga in run.text:
                            run.text = run.text.replace(palavra_antiga, palavra_nova)
                            run.bold = True  # Mantém o negrito
                        run.font.name = 'Arial'
                else:
                    for run in paragraph.runs:
                        if palavra_antiga in run.text:
                            run.text = run.text.replace(palavra_antiga, palavra_nova)
                        run.font.name = 'Arial'
        return paragraph

    def substituir_texto_tabela(self, cell, substituicoes):
        # Substituir o texto nas células da tabela
        for paragraph in cell.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)  # Chamando a função que substitui no parágrafo
        # Centralizar verticalmente o conteúdo da célula
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        return cell

    def validar_cpf(self, cpf: str) -> bool:
            """Valida um CPF informado com pontos e traço."""

            # Remover caracteres não numéricos (pontos e traço)
            cpf = re.sub(r"[\.\-]", "", cpf)

            # Verificar se tem exatamente 11 dígitos
            if not cpf.isdigit() or len(cpf) != 11:
                return False

            # Verificar se todos os dígitos são iguais (ex: 111.111.111-11 é inválido)
            if cpf == cpf[0] * 11:
                return False

            # Cálculo do primeiro dígito verificador
            soma = sum(int(cpf[i]) * (10 - i) for i in range(9))
            digito1 = (soma * 10 % 11) % 10

            # Cálculo do segundo dígito verificador
            soma = sum(int(cpf[i]) * (11 - i) for i in range(10))
            digito2 = (soma * 10 % 11) % 10

            # Verifica se os dígitos calculados são iguais aos do CPF informado
            return digito1 == int(cpf[9]) and digito2 == int(cpf[10])
        
    # Funções para TECNICO de O&M
    def nr10_sup_oem(self, substituicoes):

        # Receber os dados
        substituicoes =  {
            'NOMEFUNCIONARIO': self.entr_funcionario.get(),
            'DIAANUENCIA': datetime.today().strftime("%d"),
            'MESANUENCIA': datetime.today().strftime("%B"),
            'ANOANUENCIA': datetime.today().strftime("%Y"),
            'CPFFUNCIONARIO': self.entr_CPF.get()
        }             
   
        # Abra o documento existente
        doc = Document(os.path.join(caminho_base,'sup_oem', 'nr10_sup_oem.docx'))

        # Substitua as palavras específicas nos parágrafos
        for paragraph in doc.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)

        # Substitua as palavras específicas nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell = self.substituir_texto_tabela(cell, substituicoes)

        arq_doc_nr10 = os.path.join(caminho_base, 'sup_oem', 'NR10_temp.docx')
        arq_pdf_nr10 = os.path.join(self.dir_pasta_selecionada, f'NR10_{self.entr_apelido.get()}_{data_hoje}_SUP_OM.pdf')

        # Salve o documento editado
        doc.save(arq_doc_nr10)

        # Converta o arquivo Word para PDF usando comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(os.path.abspath(arq_doc_nr10))
        doc.SaveAs(os.path.abspath(arq_pdf_nr10), FileFormat=17)
        doc.Close()
        word.Quit()  
    
    def nr10_sep_sup_oem(self, substituicoes):

        # Receber os dados
        substituicoes =  {
            'NOMEFUNCIONARIO': self.entr_funcionario.get(),
            'DIAANUENCIA': datetime.today().strftime("%d"),
            'MESANUENCIA': datetime.today().strftime("%B"),
            'ANOANUENCIA': datetime.today().strftime("%Y"),
            'CPFFUNCIONARIO': self.entr_CPF.get()
        }    

        # Abra o documento existente
        doc = Document(os.path.join(caminho_base,'sup_oem', 'nr10_sep_sup_oem.docx'))

        # Substitua as palavras específicas nos parágrafos
        for paragraph in doc.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)

        # Substitua as palavras específicas nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell = self.substituir_texto_tabela(cell, substituicoes)
   
        arq_doc_nr10_sep = os.path.join(caminho_base,'sup_oem', 'NR10_SEP_temp.docx')
        arq_pdf_nr10_sep = os.path.join(self.dir_pasta_selecionada, f'NR10_SEP_{self.entr_apelido.get()}_{data_hoje}_SUP_OM.pdf')

        # Salve o documento editado
        doc.save(arq_doc_nr10_sep)

        # Converta o arquivo Word para PDF usando comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(os.path.abspath(arq_doc_nr10_sep))
        doc.SaveAs(os.path.abspath(arq_pdf_nr10_sep), FileFormat=17)
        doc.Close()
        word.Quit()        

    def nr12_sup_oem(self, substituicoes):

        # Receber os dados
        substituicoes =  {
            'NOMEFUNCIONARIO': self.entr_funcionario.get(),
            'DIAANUENCIA': datetime.today().strftime("%d"),
            'MESANUENCIA': datetime.today().strftime("%B"),
            'ANOANUENCIA': datetime.today().strftime("%Y"),
            'CPFFUNCIONARIO': self.entr_CPF.get()
        }    

        # Abra o documento existente
        doc = Document(os.path.join(caminho_base,'sup_oem', 'nr12_sup_oem.docx'))

        # Substitua as palavras específicas nos parágrafos
        for paragraph in doc.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)

        # Substitua as palavras específicas nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell = self.substituir_texto_tabela(cell, substituicoes)    
    
        arq_doc_nr12 = os.path.join(caminho_base,'sup_oem', 'NR12_temp.docx')
        arq_pdf_nr12 = os.path.join(self.dir_pasta_selecionada, f'NR12_{self.entr_apelido.get()}_{data_hoje}_SUP_OM.pdf')

        # Salve o documento editado
        doc.save(arq_doc_nr12)

        # Converta o arquivo Word para PDF usando comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(os.path.abspath(arq_doc_nr12))
        doc.SaveAs(os.path.abspath(arq_pdf_nr12), FileFormat=17)
        doc.Close()
        word.Quit()        

    def nr33_sup_oem(self, substituicoes):

        # Receber os dados
        substituicoes =  {
            'NOMEFUNCIONARIO': self.entr_funcionario.get(),
            'DIAANUENCIA': datetime.today().strftime("%d"),
            'MESANUENCIA': datetime.today().strftime("%B"),
            'ANOANUENCIA': datetime.today().strftime("%Y"),
            'CPFFUNCIONARIO': self.entr_CPF.get()
        }    

        # Abra o documento existente
        doc = Document(os.path.join(caminho_base,'sup_oem', 'nr33_sup_oem.docx'))

        # Substitua as palavras específicas nos parágrafos
        for paragraph in doc.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)

        # Substitua as palavras específicas nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell = self.substituir_texto_tabela(cell, substituicoes)   
    
        arq_doc_nr33 = os.path.join(caminho_base,'sup_oem', 'NR33_temp.docx')
        arq_pdf_nr33 = os.path.join(self.dir_pasta_selecionada, f'NR33_{self.entr_apelido.get()}_{data_hoje}_SUP_OM.pdf')

        # Salve o documento editado
        doc.save(arq_doc_nr33)

        # Converta o arquivo Word para PDF usando comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(os.path.abspath(arq_doc_nr33))
        doc.SaveAs(os.path.abspath(arq_pdf_nr33), FileFormat=17)
        doc.Close()
        word.Quit()        

    def nr35_sup_oem(self, substituicoes):

        # Receber os dados
        substituicoes =  {
            'NOMEFUNCIONARIO': self.entr_funcionario.get(),
            'DIAANUENCIA': datetime.today().strftime("%d"),
            'MESANUENCIA': datetime.today().strftime("%B"),
            'ANOANUENCIA': datetime.today().strftime("%Y"),
            'CPFFUNCIONARIO': self.entr_CPF.get()
        }    

        # Abra o documento existente
        doc = Document(os.path.join(caminho_base,'sup_oem', 'nr35_sup_oem.docx'))

        # Substitua as palavras específicas nos parágrafos
        for paragraph in doc.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)

        # Substitua as palavras específicas nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell = self.substituir_texto_tabela(cell, substituicoes)     
    
        arq_doc_nr35 = os.path.join(caminho_base,'sup_oem', 'NR35_temp.docx')
        arq_pdf_nr35 = os.path.join(self.dir_pasta_selecionada, f'NR35_{self.entr_apelido.get()}_{data_hoje}_SUP_OM.pdf')

        # Salve o documento editado
        doc.save(arq_doc_nr35)

        # Converta o arquivo Word para PDF usando comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(os.path.abspath(arq_doc_nr35))
        doc.SaveAs(os.path.abspath(arq_pdf_nr35), FileFormat=17)
        doc.Close()
        word.Quit()        
        
    def verificar_checkbuttons(self):      

        cpf_digitado = self.entr_CPF.get()
        texto_pastaselecionada = self.lbl_pastaselecionada.cget("text")

        if texto_pastaselecionada == "Selecione a pasta":

            messagebox.showinfo("Alerta!", "Selecione a pasta para salvar o(s) documento(s)!") 

        else:

            if self.validar_cpf(cpf_digitado):             

                notification.notify(
                    title="Aviso",
                    message="O processo está em execução, acompanhe o(s) documento(s) gerados na pasta selecionada...",
                    timeout=10  # Tempo que a notificação ficará visível (segundos)
                )

                var_check = 1
                
                # Check Tecnicos O&M
                if self.var_nr10_sup_oem.get():                
                    self.nr10_sup_oem(self.dados)  
                    notification.notify(
                        title="Aviso",
                        message="NR10 Gerada com Sucesso!",
                        timeout=5  # Tempo que a notificação ficará visível (segundos)
                    )              

                if self.var_nr10_sep_sup_oem.get():                
                    self.nr10_sep_sup_oem(self.dados)                
                    notification.notify(
                        title="Aviso",
                        message="NR10 SEP Gerada com Sucesso!",
                        timeout=5  # Tempo que a notificação ficará visível (segundos)
                    )    

                if self.var_nr12_sup_oem.get():                
                    self.nr12_sup_oem(self.dados)  
                    notification.notify(
                        title="Aviso",
                        message="NR12 Gerada com Sucesso!",
                        timeout=5  # Tempo que a notificação ficará visível (segundos)
                    )                                     

                if self.var_nr33_sup_oem.get():                
                    self.nr33_sup_oem(self.dados)               
                    notification.notify(
                        title="Aviso",
                        message="NR33 Gerada com Sucesso!",
                        timeout=5  # Tempo que a notificação ficará visível (segundos)
                    )    

                if self.var_nr35_sup_oem.get():                
                    self.nr35_sup_oem(self.dados)
                    notification.notify(
                        title="Aviso",
                        message="NR35 Gerada com Sucesso!",
                        timeout=5  # Tempo que a notificação ficará visível (segundos)
                    )                        

                # IF para sair da interção do LOOP
                if var_check == 2:
                    return
                
                else:
                    
                    # Notificação final
                    notification.notify(
                        title="Concluído",
                        message="O processo foi finalizado!",
                        timeout=10
                    )

                    messagebox.showinfo("Alerta!", "Documento(s) Salvo(s) com Sucesso!") 

                    return
                
            else:
                messagebox.showinfo("Alerta!", "CPF Inválido!") 
        
# Criar janela do Tkinter
if __name__ == "__main__":
    root = tk.Tk()
    app = genAnuencias(root)
    root.mainloop()