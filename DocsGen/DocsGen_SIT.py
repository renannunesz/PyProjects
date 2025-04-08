import os
import tkinter as tk
import comtypes.client
import re
from plyer import notification
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from ttkbootstrap import Style
from PIL import Image, ImageTk
from datetime import datetime

# Obter a data de hoje
hoje = datetime.today()
# Formatar a data como DDMMAAAA
data_hoje = hoje.strftime("%d%m%Y")

# Obtém o diretório onde o script está localizado
caminho_base = os.path.dirname(os.path.abspath(__file__))

class genSIT:
    def __init__(self, root):
        # Crie a interface gráfica
        style = Style(theme='darkly')
        self.root = root
        root = style.master
        root.title("Gerador de Anuências")   #bkg_sit
        frame = ttk.Frame(root)
        frame['padding'] = (10, 10, 10, 10)
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        self.dir_pasta_selecionada = ""

        # Carregar a imagem
        self.bg_image = Image.open(r"C:\PyProjects\DocsGen\bkg_sit.png")
        self.bg_photo = ImageTk.PhotoImage(self.bg_image)
        
        # Inserir a imagem no topo do formulário
        ttk.Label(frame, image=self.bg_photo).grid(row=0, column=0, columnspan=6, pady=(0, 20))

        #LabelFrame dentro do Frame - Dados Documento
        lblframe_dados_documentos = ttk.LabelFrame(frame, text="Dados Documento:", padding=10)
        lblframe_dados_documentos.grid(row=1, column=0, columnspan=6, sticky="ew", padx=10, pady=5)   

        # Labels
        ttk.Label(lblframe_dados_documentos, text="Funcionário: ").grid(row=0, column=0, sticky=tk.W)
        ttk.Label(lblframe_dados_documentos, text="Iniciais: ").grid(row=1, column=0, sticky=tk.W)
        ttk.Label(lblframe_dados_documentos, text="Data: ").grid(row=2, column=0, sticky=tk.W)
        ttk.Label(lblframe_dados_documentos, text="ID Curso/ID Aluno: ").grid(row=3, column=0, sticky=tk.W)

        # Entry (input)
        self.entr_funcionario = ttk.Entry(lblframe_dados_documentos, width=60)
        self.entr_funcionario.grid(row=0, column=1, columnspan=6, sticky=tk.W, pady=5)        
        self.entr_iniciais = ttk.Entry(lblframe_dados_documentos)
        self.entr_iniciais.grid(row=1, column=1, sticky=tk.W, pady=5)
        self.entr_data = ttk.Entry(lblframe_dados_documentos)
        self.entr_data.grid(row=2, column=1, sticky=tk.W, pady=5)
        ttk.Label(lblframe_dados_documentos, text=" Formato: DD/MM/AAAA").grid(row=2, column=2, sticky=tk.W)
        self.entr_id = ttk.Entry(lblframe_dados_documentos)
        self.entr_id.grid(row=3, column=1, sticky=tk.W, pady=5)
        
        #LabelFrame dentro do Frame - Opções
        lblframe_opcoes = ttk.LabelFrame(frame, text="Opções: ", padding=10)
        lblframe_opcoes.grid(row=3, column=0, columnspan=6, sticky="ew", padx=10, pady=5)                

        # Btn  
        ttk.Label(lblframe_opcoes, text="Salvar Arquivo em: ").grid(row=0, column=0, sticky=tk.W)      
        self.btn_selecionar_pasta = ttk.Button(lblframe_opcoes, text="...", command=self.selecionar_pasta).grid(row=0, column=1, sticky=tk.W, padx=(0, 10))   
        self.lbl_pastaselecionada = tk.Label(lblframe_opcoes, text="Selecione a pasta", wraplength=350)
        self.lbl_pastaselecionada.grid(row=0, column=2, columnspan=4, sticky=tk.W, padx=(0,10))

        # Botão
        self.btn_gerar = ttk.Button(frame, text="Gerar SIT", command=self.gerarDoc).grid(row=9, column=0, sticky=tk.W, padx=(0, 10), pady=10)
        self.lbl_avisoGeracao = ttk.Label(frame, text="Aguarde, os documento(s) estão sendo gerado(s)... ")

        # Defina os dados
        self.dados = {}       

    def selecionar_pasta(self):
        pasta_selecionada = filedialog.askdirectory()
        if pasta_selecionada:
            self.lbl_pastaselecionada.config(text=pasta_selecionada)   
            self.dir_pasta_selecionada = pasta_selecionada  

    def substituir_texto(self, paragraph, substituicoes):
        # Substituir o texto nos runs do parágrafo
        for palavra_antiga, palavra_nova in substituicoes.items():
            if palavra_antiga in paragraph.text:
                if palavra_antiga == 'NOMEFUNCIONARIO':
                    for run in paragraph.runs:
                        if palavra_antiga in run.text:
                            run.text = run.text.replace(palavra_antiga, palavra_nova)
                            run.bold = True  # Mantém o negrito
                        run.font.name = 'Arial'
                else:
                    for run in paragraph.runs:
                        if palavra_antiga in run.text:
                            run.text = run.text.replace(palavra_antiga, palavra_nova)
                        run.font.name = 'Arial'
        return paragraph

    def substituir_texto_tabela(self, cell, substituicoes):
        # Substituir o texto nas células da tabela
        for paragraph in cell.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)  # Chamando a função que substitui no parágrafo
        # Centralizar verticalmente o conteúdo da célula
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        return cell

    def gerar_sit(self, substituicoes):

        # Receber os dados
        substituicoes =  {
            'NOMEFUNCIONARIO': self.entr_funcionario.get(),
            'DATASIT': self.entr_data.get(),
            'IDCURSOFUNCIONARIO': self.entr_id.get()
        }             

        dataCurso = self.entr_data.get().replace("/","")
   
        # Abra o documento existente
        doc = Document(os.path.join(caminho_base,'cert_sit', 'cert_sit.docx'))

        # Substitua as palavras específicas nos parágrafos
        for paragraph in doc.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)

        # Substitua as palavras específicas nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell = self.substituir_texto_tabela(cell, substituicoes)

        arq_doc_sit = os.path.join(caminho_base, 'cert_sit', 'cert_sit_temp.docx')
        arq_pdf_sit = os.path.join(self.dir_pasta_selecionada, f'CERT_SIT_PT_{self.entr_iniciais.get()}_{dataCurso}.pdf')

        # Salve o documento editado
        doc.save(arq_doc_sit)

        # Converta o arquivo Word para PDF usando comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(os.path.abspath(arq_doc_sit))
        doc.SaveAs(os.path.abspath(arq_pdf_sit), FileFormat=17)
        doc.Close()
        word.Quit()

    def gerarDoc(self):      

        texto_pastaselecionada = self.lbl_pastaselecionada.cget("text")

        if texto_pastaselecionada == "Selecione a pasta":

            messagebox.showinfo("Alerta!", "Selecione a pasta para salvar o(s) documento(s)!") 

        else:    

            notification.notify(
                title="Aviso",
                message="O processo está em execução, acompanhe o(s) documento(s) gerados na pasta selecionada...",
                timeout=10  # Tempo que a notificação ficará visível (segundos)
            )                   

            self.gerar_sit(self.dados)
            notification.notify(
                title="Aviso",
                message="SIT Gerado com Sucesso!",
                timeout=5  # Tempo que a notificação ficará visível (segundos)
            )
                            
            # Notificação final
            notification.notify(
                title="Concluído",
                message="O processo foi finalizado!",
                timeout=10
            )

            messagebox.showinfo("Alerta!", "Documento(s) Salvo(s) com Sucesso!") 

            return
                      
# Criar janela do Tkinter
if __name__ == "__main__":
    root = tk.Tk()
    app = genSIT(root)
    root.mainloop()