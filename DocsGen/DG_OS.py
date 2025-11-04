import os
import tkinter as tk
import comtypes.client
import re
import locale
from plyer import notification
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from ttkbootstrap import Style
from PIL import Image, ImageTk
from datetime import datetime

# Seta a linguagem para Portugues do brasil.
locale.setlocale(locale.LC_TIME, "Portuguese_Brazil.1252")
# Obter a data de hoje
hoje = datetime.today()
# Formatar a data como DDMMAAAA
data_hoje = hoje.strftime("%d%m%Y")
# Retorna o diretório onde o script está localizado
caminho_base = os.path.dirname(os.path.abspath(__file__))

# Lista de meses em português
meses = {
    1: "Janeiro",
    2: "Fevereiro",
    3: "Março",
    4: "Abril",
    5: "Maio",
    6: "Junho",
    7: "Julho",
    8: "Agosto",
    9: "Setembro",
    10: "Outubro",
    11: "Novembro",
    12: "Dezembro"
}

class genOrdenServico:
    def __init__(self, root):
        # Crie a interface gráfica
        style = Style(theme='darkly')
        self.root = root
        root = style.master
        root.title("DocsGen - Ordens de Serviço")
        frame = ttk.Frame(root)
        frame['padding'] = (10, 10, 10, 10)
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        self.dir_pasta_selecionada = ""

        # Carregar a imagem
        self.bg_image = Image.open(r"C:\PyProjects\DocsGen\bkg_osst.png")
        self.bg_photo = ImageTk.PhotoImage(self.bg_image)
        
        # Inserir a imagem no topo do formulário
        ttk.Label(frame, image=self.bg_photo).grid(row=0, column=0, columnspan=6, pady=(0, 20))

        #LabelFrame dentro do Frame - Dados Documento
        lblframe_dados_documentos = ttk.LabelFrame(frame, text="Dados Documento:", padding=10)
        lblframe_dados_documentos.grid(row=6, column=0, columnspan=6, sticky="ew", padx=10, pady=5)     

        # Labels
        ttk.Label(lblframe_dados_documentos, text="Funcionário:").grid(row=1, column=0, sticky=tk.W)
        self.entr_funcionario = ttk.Entry(lblframe_dados_documentos, width=40) #width=50
        self.entr_funcionario.grid(row=1, column=1, sticky=tk.W, pady=10)

        ttk.Label(lblframe_dados_documentos, text=" CPF: ").grid(row=1, column=2, sticky=tk.W)
        self.entr_CPF = ttk.Entry(lblframe_dados_documentos, width=40)
        self.entr_CPF.grid(row=1, column=3, sticky=tk.W, pady=10)

        ttk.Label(lblframe_dados_documentos, text="Iniciais: ").grid(row=1, column=4, sticky=tk.W)
        self.entr_apelido = ttk.Entry(lblframe_dados_documentos, width=40)
        self.entr_apelido.grid(row=1, column=5, sticky=tk.W, pady=10)

        ttk.Label(lblframe_dados_documentos, text="GHE: ").grid(row=2, column=0, sticky=tk.W)
        self.cbbx_ghe = ttk.Combobox(lblframe_dados_documentos, values=["01", "02", "03", "04", "05"], width=40)
        self.cbbx_ghe.grid(row=2, column=1, sticky=tk.W, pady=10)
        self.cbbx_ghe.bind("<<ComboboxSelected>>", self.update_funcao)

        ttk.Label(lblframe_dados_documentos, text=" Função: ").grid(row=2, column=2, sticky=tk.W)
        self.cbbx_funcao = ttk.Combobox(lblframe_dados_documentos, value=["(Selecione o GHE)"], width=40)
        self.cbbx_funcao.grid(row=2, column=3, sticky=tk.W, pady=10)  

        ttk.Label(lblframe_dados_documentos, text="Obs: Os EPIs são definidos de acordo com o GHE selecionado. ").grid(row=2, column=5, columnspan=4 ,sticky=tk.W)

        ttk.Label(lblframe_dados_documentos, text="HSE Responsável: ").grid(row=3, column=0, sticky=tk.W)
        self.cbbx_hse = ttk.Combobox(lblframe_dados_documentos, values=["Leonardo Silverio", "Manoel Jefete"])
        self.cbbx_hse.grid(row=3, column=1, sticky=tk.W, pady=10)

        #LabelFrame dentro do Frame - Riscos Físicos
        lblframe_riscos_fisicos = ttk.LabelFrame(frame, text="Riscos Físicos:", padding=10)
        lblframe_riscos_fisicos.grid(row=7, column=0, columnspan=6, sticky="ew", padx=10, pady=5)

        #Variáveis para os chechbuttons riscos fisicos
        self.var_ckbx_fisc_1 = tk.BooleanVar(value=False)
        self.var_ckbx_fisc_2 = tk.BooleanVar(value=False)
        self.var_ckbx_fisc_3 = tk.BooleanVar(value=False)
        self.var_ckbx_fisc_4 = tk.BooleanVar(value=False)

        # Opcoes para garantir que fique visível
        self.ckbx_fisc_1 = ttk.Checkbutton(lblframe_riscos_fisicos, text=" • Radiação não ionizante", variable=self.var_ckbx_fisc_1)
        self.ckbx_fisc_1.grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.ckbx_fisc_2 = ttk.Checkbutton(lblframe_riscos_fisicos, text=" • Ruído contínuo ou intermitente", variable=self.var_ckbx_fisc_2)
        self.ckbx_fisc_2.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        self.ckbx_fisc_3 = ttk.Checkbutton(lblframe_riscos_fisicos, text=" • Temperaturas Anormais (calor)", variable=self.var_ckbx_fisc_3)
        self.ckbx_fisc_3.grid(row=0, column=2, sticky=tk.W, padx=5, pady=5)
        self.ckbx_fisc_4 = ttk.Checkbutton(lblframe_riscos_fisicos, text=" • Não se Aplica", variable=self.var_ckbx_fisc_4)
        self.ckbx_fisc_4.grid(row=0, column=3, sticky=tk.W, padx=5, pady=5)        

        #LabelFrame dentro do Frame - Riscos Quimicos
        lblframe_riscos_quimicos = ttk.LabelFrame(frame, text="Riscos Químicos:", padding=10)
        lblframe_riscos_quimicos.grid(row=8, column=0, columnspan=6, sticky="ew", padx=10, pady=5)

        #Variáveis para os chechbuttons riscos quimicos
        self.var_ckbx_quim_1 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_2 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_3 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_4 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_5 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_6 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_7 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_8 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_9 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_10 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_11 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_12 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_13 = tk.BooleanVar(value=False)
        self.var_ckbx_quim_14 = tk.BooleanVar(value=False)

        # Opcoes para garantir que fique visível
        self.ckbx_quim_1 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Bário e composto solúveis como Ba", variable=self.var_ckbx_quim_1)
        self.ckbx_quim_1.grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.ckbx_quim_2 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Calcio (Ca)", variable=self.var_ckbx_quim_2)
        self.ckbx_quim_2.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        self.ckbx_quim_3 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Cobre fumo como Cu", variable=self.var_ckbx_quim_3)
        self.ckbx_quim_3.grid(row=0, column=2, sticky=tk.W, padx=5, pady=5)
        self.ckbx_quim_4 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Cromo e compostos inorgânicos como Cr metal e compostos de Cr III", variable=self.var_ckbx_quim_4)
        self.ckbx_quim_4.grid(row=0, column=3, sticky=tk.W, padx=5, pady=5)

        self.ckbx_quim_5 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Estanho compostos inorgânicos e óxidos exceto hidreto de estanho", variable=self.var_ckbx_quim_5)
        self.ckbx_quim_5.grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.ckbx_quim_6 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Ferro óxido", variable=self.var_ckbx_quim_6)
        self.ckbx_quim_6.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        self.ckbx_quim_7 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Manganês e compostos inorgânicos como Mn", variable=self.var_ckbx_quim_7)
        self.ckbx_quim_7.grid(row=1, column=2, sticky=tk.W, padx=5, pady=5)
        self.ckbx_quim_8 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Níquel e compostos inorgânicos incluindo sulteto de níquel como Ni", variable=self.var_ckbx_quim_8)
        self.ckbx_quim_8.grid(row=1, column=3, sticky=tk.W, padx=5, pady=5)

        self.ckbx_quim_9 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Óleos e graxas", variable=self.var_ckbx_quim_9)
        self.ckbx_quim_9.grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        self.ckbx_quim_10 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Pentóxido de Vanádio como V", variable=self.var_ckbx_quim_10)
        self.ckbx_quim_10.grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        self.ckbx_quim_11 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Poeiras metálicas", variable=self.var_ckbx_quim_11)
        self.ckbx_quim_11.grid(row=2, column=2, sticky=tk.W, padx=5, pady=5)       
        self.ckbx_quim_12 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Alumínio metal e composto insolúveis", variable=self.var_ckbx_quim_12)
        self.ckbx_quim_12.grid(row=2, column=3, sticky=tk.W, padx=5, pady=5)   
        
        self.ckbx_quim_13 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Magnésio (Mg)", variable=self.var_ckbx_quim_13)
        self.ckbx_quim_13.grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)       
        self.ckbx_quim_14 = ttk.Checkbutton(lblframe_riscos_quimicos, text=" • Não se Aplica", variable=self.var_ckbx_quim_14) 
        self.ckbx_quim_14.grid(row=3, column=1, sticky=tk.W, padx=5, pady=5)                   

        #LabelFrame dentro do Frame - Riscos Ergonômicos
        lblframe_riscos_ergonomicos = ttk.LabelFrame(frame, text="Riscos Ergonômicos:", padding=10)
        lblframe_riscos_ergonomicos.grid(row=9, column=0, columnspan=6, sticky="ew", padx=10, pady=5)

        #Variáveis para os chechbuttons riscos Ergonômicos
        self.var_ckbx_ergo_1 = tk.BooleanVar(value=False)
        self.var_ckbx_ergo_2 = tk.BooleanVar(value=False)
        self.var_ckbx_ergo_3 = tk.BooleanVar(value=False)
        self.var_ckbx_ergo_4 = tk.BooleanVar(value=False)

        # Opcoes para garantir que fique visível
        self.ckbx_ergo_1 = ttk.Checkbutton(lblframe_riscos_ergonomicos, text=" • Esforço físico intenso", variable=self.var_ckbx_ergo_1)
        self.ckbx_ergo_1.grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.ckbx_ergo_2 = ttk.Checkbutton(lblframe_riscos_ergonomicos, text=" • Monotonia e repetitividade", variable=self.var_ckbx_ergo_2)
        self.ckbx_ergo_2.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        self.ckbx_ergo_3 = ttk.Checkbutton(lblframe_riscos_ergonomicos, text=" • Postura inadequada", variable=self.var_ckbx_ergo_3)
        self.ckbx_ergo_3.grid(row=0, column=2, sticky=tk.W, padx=5, pady=5)
        self.ckbx_ergo_4 = ttk.Checkbutton(lblframe_riscos_ergonomicos, text=" • Não se Aplica", variable=self.var_ckbx_ergo_4)
        self.ckbx_ergo_4.grid(row=0, column=3, sticky=tk.W, padx=5, pady=5)        

        #LabelFrame dentro do Frame - Riscos Mecânicos
        lblframe_riscos_mecanicos = ttk.LabelFrame(frame, text="Riscos de Acidente:", padding=10)
        lblframe_riscos_mecanicos.grid(row=10, column=0, columnspan=6, sticky="ew", padx=10, pady=5) 

        #Variáveis para os chechbuttons riscos mecanicos
        self.var_ckbx_mecan_1 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_2 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_3 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_4 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_5 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_6 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_7 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_8 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_9 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_10 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_11 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_12 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_13 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_14 = tk.BooleanVar(value=False)
        self.var_ckbx_mecan_15 = tk.BooleanVar(value=False)

        # Opcoes para garantir que fique visível
        self.ckbx_mecan_1 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Animais peçonhentos", variable=self.var_ckbx_mecan_1 )
        self.ckbx_mecan_1.grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_2 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Atmosfera explosiva", variable=self.var_ckbx_mecan_2 )
        self.ckbx_mecan_2.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_3 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Condução de veículos leves", variable=self.var_ckbx_mecan_3 )
        self.ckbx_mecan_3.grid(row=0, column=2, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_4 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Contato com objetos cortantes e/ou perfurocortantes", variable=self.var_ckbx_mecan_4 )
        self.ckbx_mecan_4.grid(row=0, column=3, sticky=tk.W, padx=5, pady=5)

        self.ckbx_mecan_5 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Eletricidade", variable=self.var_ckbx_mecan_5 )
        self.ckbx_mecan_5.grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_6 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Eletricidade estática ", variable=self.var_ckbx_mecan_6 )
        self.ckbx_mecan_6.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_7 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Espaço confinado", variable=self.var_ckbx_mecan_7 )
        self.ckbx_mecan_7.grid(row=1, column=2, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_8 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Prensamento de mãos", variable=self.var_ckbx_mecan_8 )
        self.ckbx_mecan_8.grid(row=1, column=3, sticky=tk.W, padx=5, pady=5)

        self.ckbx_mecan_9 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Projeções de fluídos e/ou ar comprimido", variable=self.var_ckbx_mecan_9 )
        self.ckbx_mecan_9.grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_10 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Projeções de objetos e fragmentos", variable=self.var_ckbx_mecan_10)
        self.ckbx_mecan_10.grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_11 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Queda de diferentes níveis (acima de 2 metros)", variable=self.var_ckbx_mecan_11)
        self.ckbx_mecan_11.grid(row=2, column=2, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_12 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Queda de objetos", variable=self.var_ckbx_mecan_12)
        self.ckbx_mecan_12.grid(row=2, column=3, sticky=tk.W, padx=5, pady=5)   

        self.ckbx_mecan_13 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Superfícies Aquecidas", variable=self.var_ckbx_mecan_13)
        self.ckbx_mecan_13.grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_14 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Trânsito de veículos pesados", variable=self.var_ckbx_mecan_14)
        self.ckbx_mecan_14.grid(row=3, column=1, sticky=tk.W, padx=5, pady=5)
        self.ckbx_mecan_15 = ttk.Checkbutton(lblframe_riscos_mecanicos, text=" • Não se Aplica", variable=self.var_ckbx_mecan_15)
        self.ckbx_mecan_15.grid(row=3, column=2, sticky=tk.W, padx=5, pady=5)

        #LabelFrame dentro do Frame - Opções
        lblframe_opcoes = ttk.LabelFrame(frame, text="Opções: ", padding=10)
        lblframe_opcoes.grid(row=11, column=0, columnspan=6, sticky="ew", padx=10, pady=5)

        # Btn      
        ttk.Label(lblframe_opcoes, text="Salvar Arquivo em: ").grid(row=11, column=0, sticky=tk.W)       
        self.btn_selecionar_pasta = ttk.Button(lblframe_opcoes, text="...", command=self.selecionar_pasta).grid(row=11, column=1, sticky=tk.W, padx=(0, 10))   
        self.lbl_pastaselecionada = tk.Label(lblframe_opcoes, text="(Selecione a pasta)", wraplength=350)
        self.lbl_pastaselecionada.grid(row=11, column=2, sticky=tk.W, padx=(0,10))

        # Botão
        self.btn_gerar = ttk.Button(frame, text="Gerar Ordens de Serviço", command=self.verificar_checkbuttons).grid(row=12, column=0, sticky=tk.W, padx=(0, 10), pady=10)

        # Defina os dados
        self.dados = {}   

    def selecionar_pasta(self):
        pasta_selecionada = filedialog.askdirectory()
        if pasta_selecionada:
            self.lbl_pastaselecionada.config(text=pasta_selecionada)   
            self.dir_pasta_selecionada = pasta_selecionada  

    def substituir_texto(self, paragraph, substituicoes):
        # Substituir o texto nos runs do parágrafo
        for palavra_antiga, palavra_nova in substituicoes.items():
            if palavra_antiga in paragraph.text:
                if palavra_antiga == 'NOMEFUNCIONARIO':
                    for run in paragraph.runs:
                        if palavra_antiga in run.text:
                            run.text = run.text.replace(palavra_antiga, palavra_nova)
                            run.bold = True  # Mantém o negrito
                        run.font.name = 'Arial'
                else:
                    for run in paragraph.runs:
                        if palavra_antiga in run.text:
                            run.text = run.text.replace(palavra_antiga, palavra_nova)
                        run.font.name = 'Arial'
        return paragraph

    def substituir_texto_tabela(self, cell, substituicoes):
        # Substituir o texto nas células da tabela
        for paragraph in cell.paragraphs:
            for chave, valor in substituicoes.items():
                self.substituir_texto(paragraph, substituicoes)  # Chamando a função que substitui no parágrafo                
        # Centralizar verticalmente o conteúdo da célula
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        return cell

    def validar_cpf(self, cpf: str) -> bool:
        """Valida um CPF informado com pontos e traço."""

        # Remover caracteres não numéricos (pontos e traço)
        cpf = re.sub(r"[\.\-]", "", cpf)

        # Verificar se tem exatamente 11 dígitos
        if not cpf.isdigit() or len(cpf) != 11:
            return False

        # Verificar se todos os dígitos são iguais (ex: 111.111.111-11 é inválido)
        if cpf == cpf[0] * 11:
            return False

        # Cálculo do primeiro dígito verificador
        soma = sum(int(cpf[i]) * (10 - i) for i in range(9))
        digito1 = (soma * 10 % 11) % 10

        # Cálculo do segundo dígito verificador
        soma = sum(int(cpf[i]) * (11 - i) for i in range(10))
        digito2 = (soma * 10 % 11) % 10

        # Verifica se os dígitos calculados são iguais aos do CPF informado
        return digito1 == int(cpf[9]) and digito2 == int(cpf[10])

    def update_funcao(self, event):
        gheSelecionado = self.cbbx_ghe.get()
        if gheSelecionado == "01":
                self.cbbx_funcao['values'] = [
                    "ANALISTA DE CUSTO",
                    "ANALISTA DE OPERACOES",
                    "ANALISTA DE NEGOCIOS",
                    "COMPRADOR SR",
                    "COORDENADOR DE FROTA",
                    "COORDENADOR DE QUALIDADE",
                    "ENGENHEIRO DE ATENDIMENTO AO CLIENTE",
                    "ENGENHEIRO DE TECNOLOGIA",
                    "PLANEJADOR",
                    "PROGRAMADOR",
                    "SUPERVISOR DE MATERIAIS",
                    "SUPERVISOR DE PLANEJAMENTO", 
                    "GERENTE DE RELACOES PUBLICAS SR",
                    "CONSULTOR ADMINISTRATIVO"                   
                    ]
        elif gheSelecionado == "02":
                self.cbbx_funcao['values'] = [
                    "ANALISTA DE QUALIDADE",
                    "ANALISTA DE TREINAMENTO", 
                    "COORDENADOR DE OPERACOES", 
                    "DIRETOR (A) DE SSMA LATAM", 
                    "ENGENHEIRO DE AUTOMACAO", 
                    "ENGENHEIRO DE CONTROLE DE QUALIDADE SR", 
                    "ENGENHEIRO DE PA EOLICA", 
                    "ENGENHEIRO DE PROJETOS", 
                    "ENGENHEIRO DE QUALIDADE", 
                    "ENGENHEIRO DE QUALIDADE SR", 
                    "ENGENHEIRO DE SEGURANÇA DO TRABALHO", 
                    "ENGENHEIRO ELETRICISTA", 
                    "ENGENHEIRO ELETRICO", 
                    "ENGENHEIRO EM GESTÃO DE PROJETOS", 
                    "ENGENHEIRO MECANICO", 
                    "ENGENHEIRO SCADA", 
                    "ESPECIALISTA DE QUALIDADE", 
                    "GERENTE DE CONTROLE DE PROJETOS", 
                    "GERENTE DE PROJETOS", 
                    "GERENTE DE PROJETOS E SERVIÇOES DE MANUTENÇÃO", 
                    "GERENTE DE QUALIDADE EXC E PROJETOS", 
                    "GERENTE DE SEGURANCA DO TRABALHO", 
                    "GERENTE SITE", 
                    "SUPERVISOR O&M", 
                    "TECNICO DE SEGURANCA DO TRABALHO", 
                    "TECNICO DE SEGURANCA DO TRABALHO SR"
                ]
        elif gheSelecionado == "03":
                self.cbbx_funcao['values'] = [
                    "SUPERVISOR DE INSTALACOES",
                    "TECNICO DE SERVICOS ESPECIAIS",
                    "TECNICO O&M ESPECIALISTA",
                    "TECNICO O&M JR",
                    "TECNICO O&M LIDER",
                    "TECNICO O&M PL",
                    "TECNICO O&M SR",
                    "TECNICO DE PA EOLICA ESPECIALISTA",
                    "TECNICO DE PA EOLICA JR",
                    "TECNICO DE PA EOLICA LIDER",
                    "TECNICO DE PA EOLICA PL",
                    "TECNICO DE PA EOLICA SR",                    
                ]
        elif gheSelecionado == "04":
                self.cbbx_funcao['values'] = [
                    "ANALISTA DE DADOS",
                    "CEO DIRETOR GERAL",
                    "DIRETOR DE CONSTRUCAO",
                    "DIRETOR DE ENGENHARIA",
                    "DIRETOR DE QUALIDADE LATAM",
                    "DIRETOR DE OPERACOES",
                    "DIRETOR DE OPERACOES E MANUTENCAO",
                    "DIRETOR DE SEGURANCA AMERICA LATINA",
                    "DIRETOR FINANCEIRO",
                    "DIRETOR FINANCEIRO CONTABIL",
                    "DIRETOR REGIONAL DE MANUFATURA",
                    "GERENTE DE ATENDIMENTO AO CLIENTE",
                    "GERENTE DE ENGENHARIA",
                    "GERENTE DE MELHORIA CONTINUA",
                    "GERENTE DE OPERACOES",
                    "GERENTE DE OPERACOES E MANUTENÇÃO",
                    "GERENTE DE RECURSOS HUMANOS",
                    "GERENTE DE SUPORTE DE OPERACAO DE SERVICO",
                    "GERENTE DE TREINAMENTO",
                    "VICE PRESIDENTE DE CONSTRUCAO",
                    "VICE PRESIDENTE DE FINANCAS",
                    "VICE PRESIDENTE DE OPERACOES",
                    "VICE PRESIDENTE REGIONAL CTO",
                    "VICE PRESIDENTE, PEOPLE & CULTURE, LATAM"
                ]           
        elif gheSelecionado == "05":
                self.cbbx_funcao['values'] = [
                    "ALMOXARIFE JR",
                    "ALMOXARIFE PL",
                    "ALMOXARIFE SR",
                    "SUPERVISOR O&M"
                ] 
    def atividade_funcao(self):
        # Abrir o arquivo para leitura
        arquivo = os.path.join(caminho_base,'listaDescFuncoes.txt')
        with open(arquivo, 'r', encoding='utf-8') as file:
            # Ler o arquivo linha por linha
            for line in file:
                # Remover espaços em branco no início e fim da linha
                line = line.strip()
                # Dividir a linha em partes usando o ponto e vírgula como separador
                parts = line.split(';')
                # Exibir as partes da linha
                #print(parts[1])
                if parts[1] == self.cbbx_funcao.get():
                    return parts[2]
    
    def gerarOS(self, substituicoes):        

        # Recebe riscos fisicos
        if self.var_ckbx_fisc_1.get():
            riscoFisico1 = self.ckbx_fisc_1.cget("text") + "\n"
        else:
            riscoFisico1 = ""
        
        if self.var_ckbx_fisc_2.get():
            riscoFisico2 = self.ckbx_fisc_2.cget("text") + "\n"
        else:
            riscoFisico2 = ""
        
        if self.var_ckbx_fisc_3.get():
            riscoFisico3 = self.ckbx_fisc_3.cget("text") + "\n"
        else:
            riscoFisico3 = ""

        if self.var_ckbx_fisc_4.get():
            riscoFisico4 = self.ckbx_fisc_4.cget("text") + "\n"
        else:
            riscoFisico4 = ""

        # Recebe riscos quimicos
        variaveisQuimicos = [
            self.var_ckbx_quim_1, 
            self.var_ckbx_quim_2, 
            self.var_ckbx_quim_3, 
            self.var_ckbx_quim_4, 
            self.var_ckbx_quim_5, 
            self.var_ckbx_quim_6, 
            self.var_ckbx_quim_7, 
            self.var_ckbx_quim_8, 
            self.var_ckbx_quim_9, 
            self.var_ckbx_quim_10, 
            self.var_ckbx_quim_11, 
            self.var_ckbx_quim_12,
            self.var_ckbx_quim_13,
            self.var_ckbx_quim_14
            ]
        
        checkboxQuimicos = [
            self.ckbx_quim_1,
            self.ckbx_quim_2,
            self.ckbx_quim_3,
            self.ckbx_quim_4,
            self.ckbx_quim_5,
            self.ckbx_quim_6,
            self.ckbx_quim_7,
            self.ckbx_quim_8,
            self.ckbx_quim_9,
            self.ckbx_quim_10,
            self.ckbx_quim_11,
            self.ckbx_quim_12,
            self.ckbx_quim_13,
            self.ckbx_quim_14
            ]
        
        riscosQuimicos = ""
        for varQuim, ckbQuim in zip(variaveisQuimicos, checkboxQuimicos):
            if varQuim.get():
                riscosQuimicos += ckbQuim.cget("text") + "\n"

        # Recebe riscos Ergonômicos
        variaveisErgonomicos = [
            self.var_ckbx_ergo_1, 
            self.var_ckbx_ergo_2, 
            self.var_ckbx_ergo_3, 
            self.var_ckbx_ergo_4
            ]
        
        checkboxErgonomicos = [
            self.ckbx_ergo_1, 
            self.ckbx_ergo_2, 
            self.ckbx_ergo_3, 
            self.ckbx_ergo_4
            ]
        
        riscosErgonomicos = ""
        for varErgo, ckbErgo in zip(variaveisErgonomicos, checkboxErgonomicos):
            if varErgo.get():
                riscosErgonomicos += ckbErgo.cget("text") + "\n"

        # Recebe riscos Mecanicos
        variaveisMecanicos = [
            self.var_ckbx_mecan_1, 
            self.var_ckbx_mecan_2, 
            self.var_ckbx_mecan_3, 
            self.var_ckbx_mecan_4, 
            self.var_ckbx_mecan_5, 
            self.var_ckbx_mecan_6, 
            self.var_ckbx_mecan_7, 
            self.var_ckbx_mecan_8, 
            self.var_ckbx_mecan_9, 
            self.var_ckbx_mecan_10, 
            self.var_ckbx_mecan_11, 
            self.var_ckbx_mecan_12, 
            self.var_ckbx_mecan_13, 
            self.var_ckbx_mecan_14, 
            self.var_ckbx_mecan_15
            ]
        
        checkboxMecanicos = [
            self.ckbx_mecan_1, 
            self.ckbx_mecan_2, 
            self.ckbx_mecan_3, 
            self.ckbx_mecan_4, 
            self.ckbx_mecan_5, 
            self.ckbx_mecan_6, 
            self.ckbx_mecan_7, 
            self.ckbx_mecan_8, 
            self.ckbx_mecan_9, 
            self.ckbx_mecan_10, 
            self.ckbx_mecan_11, 
            self.ckbx_mecan_12, 
            self.ckbx_mecan_13, 
            self.ckbx_mecan_14, 
            self.ckbx_mecan_15
            ]
        
        riscosMecanicos = ""
        for varMeca, ckbMeca in zip(variaveisMecanicos, checkboxMecanicos):
            if varMeca.get():
                riscosMecanicos += ckbMeca.cget("text") + "\n"      

        hseNome = self.cbbx_hse.get()

        if hseNome == "Leonardo Silverio":
            nomeHSE = "LEONARDO SILVERIO FERREIRA"
            registroHSE = "MTE/RN: 1360"
            funcaoHSE = "Técnico(a) de Segurança do Trabalho"
        elif hseNome == "Manoel Jefete":
            nomeHSE = "MANOEL JEFETE DA SILVA TENONIO"
            registroHSE = "MTE/RN: 1805"
            funcaoHSE = "Técnico(a) de Segurança do Trabalho"
        
        #["Bruna Petroni", "Leonardo Silverio", "Manoel Jefete"])

        # Receber os dados
        substituicoes =  {
            'NOMEFUNCIONARIO': self.entr_funcionario.get(),
            'CPFFUNCIONARIO': self.entr_CPF.get(),
            'FUNCFUNCIONARIO': self.cbbx_funcao.get(),
            'ATVFUNCIONARIO': self.atividade_funcao(),
            'DIAOS': datetime.today().strftime("%d"),
            'MESOS': meses[datetime.today().month],
            'ANOOS': datetime.today().strftime("%Y"),
            'NOMEHSE': nomeHSE,
            'REGISTROHSE': registroHSE,
            'FUNCAOHSE': funcaoHSE
        }        

        # Recebe tipo de GHE
        tipoGHE = self.cbbx_ghe.get()

        if tipoGHE == "01":
            documentoModelo = 'osst_ghe01.docx'
        elif tipoGHE == "02":
            documentoModelo = 'osst_ghe02.docx'
        elif tipoGHE == "03":
            documentoModelo = 'osst_ghe03.docx'
        elif tipoGHE == "04":
            documentoModelo = 'osst_ghe04.docx'
        else:
            print("Outros")

        # Abra o documento existente
        doc = Document(os.path.join(caminho_base,'osst_mod', documentoModelo))

        # Substitua as palavras específicas nos parágrafos
        for paragraph in doc.paragraphs:
            paragraph = self.substituir_texto(paragraph, substituicoes)

        # Substitua as palavras específicas nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell = self.substituir_texto_tabela(cell, substituicoes)

        arq_osst = os.path.join(caminho_base, 'osst_mod', 'osst_temp.docx')
        arq_pdf_osst = os.path.join(self.dir_pasta_selecionada, f'OSST_{self.entr_apelido.get()}_{data_hoje}.pdf')

        # Acessar a primeira tabela (índice 0)
        tabelaRiscosFisicos = doc.tables[1]
        tabelaRiscosQuimicos = doc.tables[2]
        tabelaRiscosErgonomicos = doc.tables[3]
        tabelaRiscosMecanicos = doc.tables[4]
        # Definir a linha e a coluna da célula desejada
        #cell(linha, coluna) (índice começa em 0)

        # Acessar a célula e modificar o conteúdo
        tabelaRiscosFisicos.cell(0, 1).text = riscoFisico1 + riscoFisico2 + riscoFisico3 + riscoFisico4
        tabelaRiscosQuimicos.cell(0, 1).text = riscosQuimicos
        tabelaRiscosErgonomicos.cell(0, 1).text = riscosErgonomicos
        tabelaRiscosMecanicos.cell(0, 1).text = riscosMecanicos

        # Salve o documento editado
        doc.save(arq_osst)

        # Converta o arquivo Word para PDF usando comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(os.path.abspath(arq_osst))
        doc.SaveAs(os.path.abspath(arq_pdf_osst), FileFormat=17)
        doc.Close()
        word.Quit()

    def verificar_checkbuttons(self):
        cpf_digitado = self.entr_CPF.get()
        texto_pastaselecionada = self.lbl_pastaselecionada.cget("text")
        if texto_pastaselecionada == "(Selecione a pasta)":
            messagebox.showinfo("Alerta!", "Selecione a pasta para salvar o(s) documento(s)!") 
        else:
            if self.validar_cpf(cpf_digitado):

                # Notificação Inicio
                notification.notify(
                    title="Aviso",
                    message="O processo está em execução, acompanhe o(s) documento(s) gerados na pasta selecionada...",
                    timeout=10  # Tempo que a notificação ficará visível (segundos)
                )

                self.gerarOS(self.dados)

                # Notificação Finalização
                notification.notify(
                    title="Concluído",
                    message="O processo foi finalizado!",
                    timeout=10
                )
                messagebox.showinfo("Alerta!", "Arquivos Salvos com Sucesso!") 
                return
            
            else:
                messagebox.showinfo("Alerta!", "CPF Inválido!") 
        
# Criar janela do Tkinter
if __name__ == "__main__":
    root = tk.Tk()
    app = genOrdenServico(root)
    root.mainloop()